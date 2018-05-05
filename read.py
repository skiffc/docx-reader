from docx import Document
from docx.shared import Inches
from iter_block_items import iter_block_items 
from docx.table import Table 
from docx.text.paragraph import Paragraph 


document = Document('target.docx')
#print document.paragraphs
#
#for p in document.paragraphs:
#    print p.text
#
#for t in document.tables:
#    for c in t.row_cells(1):
#        print c.text

for b in iter_block_items(document):
    print b
    if isinstance(b, Paragraph):
        print b.text
    elif isinstance(b, Table):
        print 'Table'
        for row in b.rows:
            print '  Row'
            for cell in row.cells:
                print '    Cell'
                for paragraph in cell.paragraphs:
                    print '      Text:', paragraph.text

        #for c in range(0,len(b.columns)):
        #    print 'Column:', c
        #    for t in b.row_cells(c):
        #        print ' ', c, t.text
